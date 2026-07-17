import unittest
import tempfile
from pathlib import Path
from unittest.mock import patch
from docx import Document

from verifier.company import CompanyRegistry
from verifier.config import AppConfig
from verifier.idcard import birthday_from_id, validate_cn_id
from verifier.models import Evidence, Material, WorkRecord
from verifier.normalize import duration_months, format_year_month, normalize_date
from verifier.rules import evaluate
from verifier.readers import classify_document, refine_document_type, _docx_pages, read_material
from verifier.extract import extract_material
from verifier.quality import has_red_stamp
from verifier.quality import assess_id_photo
from verifier.ocr import _extract_rapid_text
from PIL import Image, ImageDraw


def material(person, filename, kind, evidences=()):
    value = Material(person, Path(filename), kind)
    value.evidences = list(evidences)
    return value


def evidence(person, filename, kind, field, raw, normalized=None):
    return Evidence(person, filename, 1, kind, field, raw, normalized if normalized is not None else raw)


class RuleTests(unittest.TestCase):
    def test_picture_filename_is_classified_by_visual_content(self):
        class FakeOCR:
            last_confidence = None
            calls = 0
            def recognize(self, image):
                self.calls += 1
                return ""

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "袁颖潞_picture-0.png"
            Image.new("RGB", (800, 1000), (240, 220, 210)).save(path)
            fake = FakeOCR()
            with patch("verifier.readers.assess_id_photo", return_value=[]):
                result = read_material("袁颖潞", path, AppConfig(), fake)
            self.assertEqual(result.document_type, "证件照")
            self.assertEqual(fake.calls, 0)
            self.assertFalse(result.errors)

    def test_misleading_filename_cannot_override_document_content(self):
        class FakeOCR:
            last_confidence = 0.99
            def recognize(self, image):
                return "企业名称 闽清金鑫物业管理有限公司 统一社会信用代码 91350124565353903Y 登记状态 存续 经营范围 物业管理"

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "身份证.png"
            Image.new("RGB", (1000, 700), "white").save(path)
            with patch("verifier.readers.assess_id_photo", return_value=["不是证件照"]):
                result = read_material("测试", path, AppConfig(), FakeOCR())
            self.assertEqual(result.document_type, "企业信息截图")

    def test_form_irregular_birth_and_graduation_dates_are_not_missing(self):
        text = (
            "福建省职业技能等级认定申报表\n出生年月\n联系信息\n2004 年 10 月\n"
            "最高学历 高职 毕业院校 福建农业职业技术学院\n毕业时间\n填写值\n2027 年 6 月 26 日\n"
        )
        form = Material("袁颖潞", Path("任意文件.pdf"), "申报表")
        form.text_pages = [text]
        evidences, _ = extract_material(form)
        fields = {e.field: e.normalized_value for e in evidences}
        self.assertEqual(fields["出生日期"], "2004-10")
        self.assertEqual(fields["毕业时间"], "2027-06-26")

    def test_docx_vertical_merged_work_label_does_not_hide_data_row(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "陈韵歆申报表.docx"
            doc = Document()
            table = doc.add_table(rows=3, cols=5)
            table.cell(0, 0).merge(table.cell(2, 0)).text = "工作经历"
            for index, value in enumerate(["何年至何年", "从事何职业", "所在单位", "证明人姓名、电话"], 1):
                table.cell(0, index).text = value
            for index, value in enumerate(["2025.4 至今", "物业电工", "闽清金鑫物业管理有限公司", "黄先生 15959179257"], 1):
                table.cell(1, index).text = value
            doc.save(path)
            form = Material("陈韵歆", path, "申报表")
            form.text_pages = _docx_pages(path)
            _, records = extract_material(form)
            self.assertEqual(len(records), 1)
            self.assertEqual(records[0].start, "2025-04")
            self.assertEqual(records[0].end, "至今")
            self.assertEqual(records[0].occupation, "物业电工")
            self.assertEqual(records[0].company, "闽清金鑫物业管理有限公司")
            self.assertEqual(records[0].witness_phone, "15959179257")

    def test_pdf_form_work_cells_split_across_lines(self):
        text = (
            "福建省职业技能等级认定申报表\n工\n作\n经\n历\n"
            "何年至何年\n从事何职业\n所在单位\n证明人姓名、电话\n"
            "2025.4 至今\n物业电工\n闽清金鑫物业管理\n有限公司\n黄先生 15959179257\n"
        )
        form = Material("陈韵歆", Path("陈韵歆申报表.pdf"), "申报表")
        form.text_pages = [text]
        _, records = extract_material(form)
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0].company, "闽清金鑫物业管理有限公司")
        self.assertEqual(records[0].occupation, "物业电工")
        self.assertEqual(records[0].end, "至今")
        self.assertEqual(records[0].witness_phone, "15959179257")

    def test_docx_work_table_preserves_rows_and_cells(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "测试申报表.docx"
            doc = Document()
            table = doc.add_table(rows=3, cols=4)
            headers = ["工作经历", "从事何职业", "所在单位", "证明人姓名、电话"]
            for index, value in enumerate(headers):
                table.rows[0].cells[index].text = value
            values = ["2023年2月至2024年6月", "物业电工", "闽清金鑫物业管理有限公司", "黄先生 15959179257"]
            for index, value in enumerate(values):
                table.rows[1].cells[index].text = value
            values = ["2024年7月至今", "物业电工", "闽清金鑫物业管理有限公司", "林先生 13800138000"]
            for index, value in enumerate(values):
                table.rows[2].cells[index].text = value
            doc.save(path)
            form = Material("测试", path, "申报表")
            form.text_pages = _docx_pages(path)
            _, records = extract_material(form)
            self.assertEqual(len(records), 2)
            self.assertEqual(records[0].company, "闽清金鑫物业管理有限公司")
            self.assertEqual(records[1].end, "至今")

    def test_all_eight_material_types_appear_in_completeness(self):
        result = evaluate("测试", [], [], [], CompanyRegistry())
        fields = {f.field: f.status for f in result.findings if f.category == "材料完整性"}
        expected = {"申报表", "证件照", "身份证", "学历证明", "工作证明", "学信网材料", "工作年限承诺书", "企业信息截图"}
        self.assertTrue(expected.issubset(fields))
        self.assertEqual(fields["工作证明"], "不适用")
        self.assertEqual(fields["企业信息截图"], "不适用")

    def test_higher_education_material_is_selected(self):
        person = "测试人员"
        high = material(person, "高中毕业证.jpg", "学历证明", [
            evidence(person, "高中毕业证.jpg", "学历证明", "学历层次", "高中"),
            evidence(person, "高中毕业证.jpg", "学历证明", "毕业院校", "示例中学"),
        ])
        bachelor = material(person, "本科毕业证.jpg", "学历证明", [
            evidence(person, "本科毕业证.jpg", "学历证明", "学历层次", "本科"),
            evidence(person, "本科毕业证.jpg", "学历证明", "毕业院校", "示例大学"),
            evidence(person, "本科毕业证.jpg", "学历证明", "毕业证编码", "1234567890"),
        ])
        evaluate(person, [high, bachelor], high.evidences + bachelor.evidences, [], CompanyRegistry())
        self.assertFalse(high.selected_as_basis)
        self.assertTrue(bachelor.selected_as_basis)

    def test_business_scope_is_matched_to_same_company(self):
        person = "测试人员"
        company_a = "甲物业管理有限公司"
        company_b = "乙机电设备有限公司"
        record = WorkRecord(person, company_a, "2024-01", "至今", None, "申报表 第1页", occupation="电工", source_type="申报表", witness_name="张三", witness_phone="13800138000")
        shot_a = material(person, "甲公司.png", "企业信息截图", [
            evidence(person, "甲公司.png", "企业信息截图", "企业名称", company_a),
            evidence(person, "甲公司.png", "企业信息截图", "经营范围", "物业管理"),
        ])
        shot_b = material(person, "乙公司.png", "企业信息截图", [
            evidence(person, "乙公司.png", "企业信息截图", "企业名称", company_b),
            evidence(person, "乙公司.png", "企业信息截图", "经营范围", "电气设备维修"),
        ])
        result = evaluate(person, [shot_a, shot_b], shot_a.evidences + shot_b.evidences, [record], CompanyRegistry())
        findings = [f for f in result.findings if f.category == "经营范围核对"]
        self.assertTrue(findings)
        self.assertEqual(findings[0].status, "人工复核")

    def test_low_confidence_key_field_requires_review(self):
        person = "测试人员"
        m = material(person, "身份证.jpg", "身份证", [
            Evidence(person, "身份证.jpg", 1, "身份证", "姓名", "测式人员", "测式人员", confidence=0.52)
        ])
        result = evaluate(person, [m], m.evidences, [], CompanyRegistry())
        self.assertTrue(any(f.category == "识别置信度" and f.field == "姓名" for f in result.findings))

    def test_form_table_work_records_are_structured(self):
        text = (
            "福建省职业技能等级认定申报表\n"
            "何年何月至何年何月\t从事何职业\t所在单位\t证明人姓名、电话\n"
            "2024年7月至今\t物业电工\t闽清金鑫物业管理有限公司\t黄先生，15959179257"
        )
        form = Material("陈俊顺", Path("陈俊顺申报表.docx"), "申报表")
        form.text_pages = [text]
        _, records = extract_material(form)
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0].company, "闽清金鑫物业管理有限公司")
        self.assertEqual(records[0].occupation, "物业电工")
        self.assertEqual(records[0].start, "2024-07")
        self.assertEqual(records[0].end, "至今")
        self.assertEqual(records[0].witness_name, "黄先生")
        self.assertEqual(records[0].witness_phone, "15959179257")

    def test_identity_form_and_school_field_extraction(self):
        identity = Material("李锴铄", Path("李锴铄身份证.pdf"), "身份证")
        identity.text_pages = ["姓名 李锴铄 性别 男 民族 汉 出生 2009年2月24日\n公民身份号码 350123200902241234"]
        identity_fields = {e.field: e.normalized_value for e in extract_material(identity)[0]}
        self.assertEqual(identity_fields["姓名"], "李锴铄")

        form = Material("李锴铄", Path("李锴铄申报表.docx"), "申报表")
        form.text_pages = ["福建省职业技能等级认定申报表\n姓名\t李锴铄\t身份证号码\t350123200902241234"]
        form_fields = {e.field: e.normalized_value for e in extract_material(form)[0]}
        self.assertEqual(form_fields["姓名"], "李锴铄")
        self.assertEqual(form_fields["身份证号"], "350123200902241234")

        diploma = Material("李锴铄", Path("李锴铄学历.jpg"), "学历证明")
        diploma.text_pages = ["学生 李锴铄，于二〇二一年九月至二〇二四年六月在本校初中部学习。\n学校：平潭翰英中学 校长：某某\n（初）毕字（24）第11380081号"]
        diploma_fields = {e.field: e.normalized_value for e in extract_material(diploma)[0]}
        self.assertEqual(diploma_fields["毕业院校"], "平潭翰英中学")

    def test_color_portrait_does_not_require_ocr(self):
        image = Image.new("RGB", (413, 579), (70, 120, 180))
        # 没有人脸时可以报告构图问题，但不能产生“未提取到文字”。
        reasons = assess_id_photo(image)
        self.assertFalse(any("文字" in reason for reason in reasons))

    def test_rapidocr_output_adapter(self):
        class Output:
            txts = ["教育部学籍在线验证报告", "学校名称 福建农业职业技术学院", ""]
            scores = [0.99, 0.91, 0.80]

        text = _extract_rapid_text(Output())
        self.assertIn("学籍在线验证报告", text)
        self.assertIn("福建农业职业技术学院", text)
        legacy = ([[[0, 0], "毕业证书", 0.98], [[0, 0], "低置信噪声", 0.1]], 0.2)
        self.assertEqual(_extract_rapid_text(legacy), "毕业证书")

    def test_junior_high_diploma_recognition(self):
        self.assertEqual(classify_document(Path("陈俊顺学历.jpg")), "学历证明")
        self.assertEqual(classify_document(Path("黄苗可毕业证书.jpg")), "学历证明")
        ocr_text = "学生 陈俊顺，于二〇二一年九月至二〇二四年七月 在本校初中部学习，学制叁年，修业期满，成绩合格，准予毕业。（初）毕字（24）第08150197号 二〇二四年七月十五日"
        self.assertEqual(refine_document_type("其他材料", ocr_text), "学历证明")
        diploma = Material("陈俊顺", Path("材料1.jpg"), refine_document_type("其他材料", ocr_text))
        diploma.text_pages = [ocr_text]
        evidences, _ = extract_material(diploma)
        level = next(e.normalized_value for e in evidences if e.field == "学历层次")
        self.assertEqual(level, "初中")

    def test_company_registry_screenshot_recognition(self):
        company_page = "企业名称 福州德鲁伊文化传媒有限公司 统一社会信用代码 91350100MABQNX33XG 经营状态 注销 成立日期 2022-06-16 登记机关 福州市市场监督管理局 经营范围 文艺创作；软件开发"
        sole_trader_page = "工商信息 名称 福州市鼓楼区美日甜甜品店 经营者 蔡艺萱 登记状态 存续 统一社会信用代码 92350102MABXEFRX6L 企业类型 个体工商户 登记机关 福州市鼓楼区市场监督管理局 经营范围 餐饮服务；食品销售"
        self.assertEqual(classify_document(Path("22.png")), "其他材料")
        self.assertEqual(refine_document_type("其他材料", company_page), "企业信息截图")
        self.assertEqual(refine_document_type("其他材料", sole_trader_page), "企业信息截图")

    def test_high_school_and_bachelor_diplomas(self):
        high_school = "普通高中毕业证书 学生彭思敏 自二〇一八年九月至二〇二一年七月在本校修业期满，成绩合格，准予毕业。学校：三合中学 毕业证号：20210827040385 二〇二一年七月十日"
        bachelor = "毕业证书 学生陈璐 于2007年1月至2010年1月在本校英语专业业余学习，修完三年制专升本科教学计划规定的全部课程，成绩合格，准予毕业。校名：漳州师范学院 证书编号：104025201005000449"
        self_study = "高等教育自学考试毕业证书 姓名庄艳华 参加人力资源管理专业本科高等教育自学考试，全部课程成绩合格，经审定，准予毕业。高等院校：福建师范大学 证书编号65359101211143622"
        for content, expected_level in ((high_school, "高中"), (bachelor, "本科"), (self_study, "本科")):
            diploma = Material("测试", Path("毕业证书.jpg"), refine_document_type("其他材料", content))
            diploma.text_pages = [content]
            evidences, _ = extract_material(diploma)
            self.assertEqual(diploma.document_type, "学历证明")
            self.assertEqual(next(e.normalized_value for e in evidences if e.field == "学历层次"), expected_level)
        self_study_diploma = Material("测试", Path("毕业证书.pdf"), "学历证明")
        self_study_diploma.text_pages = [self_study]
        evidences, _ = extract_material(self_study_diploma)
        self.assertEqual(next(e.normalized_value for e in evidences if e.field == "学历形式"), "高等教育自学考试")

    def test_chsi_subtypes_and_skill_certificate(self):
        student = "教育部学籍在线验证报告 学校名称 福建农业职业技术学院 层次 专科 学籍状态 在籍（注册学籍） 入学日期 2024年09月08日 预计毕业日期 2027年06月26日 在线验证码 ARTJPH2ZKNDQ22CK"
        degree = "中国高等教育学位在线验证报告 学位授予单位 福建师范大学 所授学位 艺术学学士学位 获学位日期 2019年06月17日 学位证书编号 1039442019003115"
        skill = "职业技能等级证书 姓名 吴海娟 职业名称 互联网营销师 工种/职业方向 平台管理员 职业技能等级 三级/高级工 证书编号 S000035001051243000362"
        self.assertEqual(refine_document_type("其他材料", student), "学信网学籍证明")
        self.assertEqual(refine_document_type("学历证明", degree), "学信网学位证明")
        self.assertEqual(refine_document_type("其他材料", skill), "职业技能等级证书")
        m = Material("吴海娟", Path("证书.jpg"), "职业技能等级证书")
        m.text_pages = [skill]
        evidences, _ = extract_material(m)
        fields = {e.field: e.normalized_value for e in evidences}
        self.assertEqual(fields["职业名称"], "互联网营销师")
        self.assertEqual(fields["职业技能等级"], "三级/高级工")
        self.assertEqual(fields["职业证书编号"], "S000035001051243000362")

    def test_chsi_student_school_expected_graduation_and_diploma_exemption(self):
        person = "蔡晴"
        report_text = (
            "教育部学籍在线验证报告 姓名 蔡晴 学校名称 福建农业职业技术学院 "
            "层次 专科 专业 现代通信技术 学制 3年 学历类别 普通高等教育 "
            "学习形式 普通全日制 分院 系所 信息工程学院 入学日期 2024年09月08日 "
            "学籍状态 在籍（注册学籍） 预计毕业日期 2027年06月26日"
        )
        chsi = Material(person, Path("蔡晴_教育部学籍在线验证报告.pdf"), "学信网学籍证明")
        chsi.text_pages = [report_text]
        chsi.evidences, _ = extract_material(chsi)
        fields = {e.field: e.normalized_value for e in chsi.evidences}
        self.assertEqual(fields["毕业院校"], "福建农业职业技术学院")
        self.assertEqual(fields["学历层次"], "大专")
        self.assertEqual(fields["学籍状态"], "在籍")
        self.assertEqual(fields["预计毕业时间"], "2027-06-26")
        self.assertEqual(fields["毕业时间"], "2027-06-26")

        form_evidences = [
            evidence(person, "申报表.pdf", "申报表", "毕业院校", "福建农业职业技术学院"),
            evidence(person, "申报表.pdf", "申报表", "毕业时间", "2027-06-26"),
        ]
        form = material(person, "申报表.pdf", "申报表", form_evidences)
        form.text_pages = ["福建省职业技能等级认定申报表 初中"]
        materials = [form, chsi, material(person, "证件照.jpg", "证件照"), material(person, "身份证.pdf", "身份证")]
        result = evaluate(person, materials, form_evidences + chsi.evidences, [], CompanyRegistry())
        findings = {(f.category, f.field): f for f in result.findings}
        self.assertEqual(findings[("学习经历核对", "毕业院校")].status, "一致")
        self.assertEqual(findings[("学习经历核对", "毕业时间")].status, "一致")
        self.assertEqual(findings[("学习经历核对", "学籍状态")].status, "在籍")
        self.assertEqual(findings[("材料完整性", "学历证明")].status, "不适用")

    def test_graduated_higher_education_still_requires_diploma(self):
        person = "已毕业人员"
        chsi = material(person, "学历在线验证报告.pdf", "学信网学历证明", [
            evidence(person, "学历在线验证报告.pdf", "学信网学历证明", "学历层次", "大专"),
            evidence(person, "学历在线验证报告.pdf", "学信网学历证明", "学籍状态", "毕业"),
            evidence(person, "学历在线验证报告.pdf", "学信网学历证明", "毕业院校", "示例职业技术学院"),
            evidence(person, "学历在线验证报告.pdf", "学信网学历证明", "毕业时间", "2024-06-30"),
        ])
        result = evaluate(person, [chsi], chsi.evidences, [], CompanyRegistry())
        self.assertTrue(any(f.category == "材料完整性" and f.field == "学历证明" and f.status == "缺少材料" for f in result.findings))

    def test_commitment_word_table_and_month_total(self):
        text = "工作年限承诺书\n姓名: 彭思敏，现申请参加 公共营养师 (职业/工种)_4__级职业技能等级认定，从事本职业或相关职业工作共5年，工作经历如下：\n2021年8月至2022年12月\t皇冠蛋糕店\t遂川县\t西点裱花师\n2023年2月至2026年5月\t可斯贝莉门店\t泉州市\t西点裱花师\n2026年6月至2026年7月\t福州市鼓楼区美日甜甜品店\t福州市\t西点裱花师\n考生签名：彭思敏"
        m = Material("彭思敏", Path("工作年限承诺书.docx"), "工作年限承诺书")
        m.text_pages = [text]
        evidences, records = extract_material(m)
        self.assertEqual(len(records), 3)
        self.assertEqual(sum(duration_months(w.start, w.end) or 0 for w in records), 59)
        fields = {e.field: e.normalized_value for e in evidences}
        self.assertEqual(fields["承诺工作年限"], "60")
        self.assertEqual(fields["承诺人签名"], "彭思敏")

    def test_work_proof_and_stamp_detection(self):
        text = "工作证明\n兹有我单位 陈俊顺，身份证号码：350181200810080357\n自 2025 年 1 月至今，在我单位从事物业电工相关行业工作。\n部门联系人：黄先生\n联系电话：15959179257\n单位（盖章）：闽清金鑫物业管理有限公司"
        m = Material("陈俊顺", Path("工作证明.docx"), "工作证明")
        m.text_pages = [text]
        evidences, records = extract_material(m)
        fields = {e.field: e.normalized_value for e in evidences}
        self.assertEqual(fields["证明人姓名"], "黄先生")
        self.assertEqual(fields["证明人电话"], "15959179257")
        self.assertEqual(fields["出具单位"], "闽清金鑫物业管理有限公司")
        self.assertEqual(records[0].end, "至今")
        image = Image.new("RGB", (1000, 1400), "white")
        ImageDraw.Draw(image).ellipse((650, 950, 930, 1230), outline=(210, 0, 0), width=35)
        self.assertTrue(has_red_stamp(image))

    def test_dates_chinese_and_duration(self):
        self.assertEqual(normalize_date("2020年3月"), "2020-03")
        self.assertEqual(normalize_date("二〇二〇年七月"), "2020-07")
        self.assertEqual(normalize_date("二零二零年七月十五日"), "2020-07-15")
        self.assertEqual(normalize_date("二O二四年七月十五日"), "2024-07-15")
        self.assertEqual(format_year_month("2018-07-07"), "2018.7")
        self.assertEqual(duration_months("2020-03", "2021-02"), 12)

    def test_form_labels_are_not_values(self):
        text = "福建省职业技能等级认定申报表\n姓名\t陈俊顺\t性别\t男\n最高学历\t初中/中职/高中\t毕业院校\n申报职业\t申报等级\t现工作单位\n工作经历\t从事何职业\t所在单位\t证明人姓名、电话\n学习经历\t初中\t高职/本科"
        m = Material("陈俊顺", Path("陈俊顺申报表.docx"), "申报表")
        m.text_pages = [text]
        evidences, _ = extract_material(m)
        fields = {e.field: e.normalized_value for e in evidences}
        self.assertEqual(fields.get("姓名"), "陈俊顺")
        self.assertNotIn("学历层次", fields)
        self.assertNotIn("企业名称", fields)
        self.assertNotIn("证明人姓名", fields)
        self.assertNotIn("从事职业", fields)

    def test_id_validation(self):
        ok, reasons = validate_cn_id("11010519491231002X")
        self.assertTrue(ok)
        self.assertFalse(reasons)
        self.assertEqual(birthday_from_id("11010519491231002X"), "1949-12-31")

    def test_company_hard_rule(self):
        registry = CompanyRegistry()
        self.assertEqual(registry.validate("腾讯")[0], "退回")
        self.assertIn("简称", registry.validate("腾讯公司")[1])
        self.assertEqual(registry.validate("深圳市腾讯计算机系统有限公司")[0], "待复核")

    def test_form_identity_mismatch(self):
        person = "张三"
        identity_evidence = [
            evidence(person, "身份证.jpg", "身份证", "姓名", "张三"),
            evidence(person, "身份证.jpg", "身份证", "身份证号", "11010519491231002X"),
            evidence(person, "身份证.jpg", "身份证", "身份证有效期至", "长期"),
        ]
        form_evidence = [
            evidence(person, "申报表.docx", "申报表", "姓名", "张山"),
            evidence(person, "申报表.docx", "申报表", "身份证号", "11010519491231002X"),
        ]
        materials = [
            material(person, "身份证.jpg", "身份证", identity_evidence),
            material(person, "申报表.docx", "申报表", form_evidence),
            material(person, "证件照.jpg", "证件照"),
            material(person, "毕业证.jpg", "学历证明"),
        ]
        result = evaluate(person, materials, identity_evidence + form_evidence, [], CompanyRegistry())
        self.assertTrue(any(f.field == "姓名" and f.status == "不一致" for f in result.findings))

    def test_two_jobs_need_commitment_and_no_overlap(self):
        person = "李四"
        identity_evidence = [
            evidence(person, "身份证.jpg", "身份证", "身份证号", "11010519491231002X"),
            evidence(person, "身份证.jpg", "身份证", "身份证有效期至", "长期"),
        ]
        form = material(person, "申报表.docx", "申报表")
        identity = material(person, "身份证.jpg", "身份证", identity_evidence)
        jobs = [
            WorkRecord(person, "北京示例科技有限公司", "2020-01", "2021-06", None, "申报表 第1页", source_type="申报表"),
            WorkRecord(person, "上海示例科技有限公司", "2021-06", "2022-12", None, "申报表 第1页", source_type="申报表"),
        ]
        materials = [form, identity, material(person, "证件照.jpg", "证件照"), material(person, "毕业证.jpg", "学历证明"), material(person, "工作证明.jpg", "工作证明"), material(person, "企业1.png", "企业信息截图"), material(person, "企业2.png", "企业信息截图")]
        result = evaluate(person, materials, identity_evidence, jobs, CompanyRegistry(), AppConfig())
        self.assertTrue(any(f.field == "工作年限承诺书" and f.status == "缺少材料" for f in result.findings))
        self.assertTrue(any(f.field == "工作经历重叠" and f.status == "不一致" for f in result.findings))

    def test_work_gap_is_allowed(self):
        person = "王五"
        jobs = [
            WorkRecord(person, "甲有限公司", "2020-01", "2020-06", None, "表", source_type="申报表"),
            WorkRecord(person, "乙有限公司", "2021-01", "2021-06", None, "表", source_type="申报表"),
        ]
        result = evaluate(person, [], [], jobs, CompanyRegistry())
        self.assertFalse(any(f.field == "工作经历重叠" for f in result.findings))


if __name__ == "__main__":
    unittest.main()
