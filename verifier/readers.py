from __future__ import annotations

import io
import re
from pathlib import Path

import fitz
from docx import Document
from PIL import Image, ImageOps

from .config import AppConfig
from .models import Material
from .ocr import LocalTesseractOCR
from .quality import assess_id_image, assess_id_photo, has_red_stamp


def classify_document(path: Path) -> str:
    n = path.stem.lower()
    rules = [
        ("申报表", ("福建省职业技能等级认定申报表", "认定申报表", "申报表", "报名表")),
        ("证件照", ("证件照", "登记照", "一寸照", "二寸照", "照片", "photo")),
        ("身份证", ("身份证", "idcard", "id_card")),
        ("学信网学籍证明", ("学籍在线验证", "教育部学籍", "学籍证明")),
        ("学信网学历证明", ("学历证书电子注册备案表", "高等教育学历认证报告", "学历在线验证")),
        ("学信网学位证明", ("学位在线验证", "学位认证报告")),
        ("学历证明", ("毕业证", "毕业证书", "学历证", "学历证明", "学历", "毕业证明", "初中毕业", "高中毕业", "中职毕业", "高职毕业", "本科毕业", "研究生毕业")),
        ("学位证", ("学位证",)),
        ("职业技能等级证书", ("职业技能等级证书", "技能等级证书")),
        ("工作年限承诺书", ("工作年限承诺", "年限承诺", "承诺函", "承诺书")),
        ("企业信息截图", ("企业信息", "工商信息", "经营范围", "企查查", "天眼查", "爱企查", "国家企业信用")),
        ("劳动合同", ("劳动合同", "聘用合同", "合同")),
        ("离职证明", ("离职", "解除劳动", "终止劳动")),
        ("工作证明", ("工作证明", "任职证明", "在职证明")),
        ("简历", ("简历", "resume", "cv")),
    ]
    for label, keys in rules:
        if any(k in n for k in keys):
            return label
    return "其他材料"


def refine_document_type(kind: str, text: str) -> str:
    compact = re.sub(r"\s", "", text)
    # 初中毕业证常不印“毕业证书”标题，以固定正文组合识别。
    junior_diploma = (
        ("初中部" in compact or "初中学习" in compact)
        and ("准予毕业" in compact or "毕字" in compact)
    )
    if junior_diploma:
        return "学历证明"
    diploma_markers = (
        "普通高中毕业证书", "高中毕业证书", "中等职业学校毕业证书",
        "中等专业学校毕业证书", "技工学校毕业证书", "高等教育自学考试毕业证书",
        "成人高等教育毕业证书", "普通高等学校毕业证书",
    )
    if any(marker in compact for marker in diploma_markers):
        return "学历证明"
    # 工商查询平台截图的文件名可能只是序号，也可能是企业全称。
    # 以工商字段组合识别，并兼容公司、企业和个体工商户。
    registry_name = any(key in compact for key in ("企业名称", "主体名称", "名称", "字号名称"))
    registry_fields = (
        "统一社会信用代码", "工商注册号", "注册号", "纳税人识别号",
        "登记状态", "经营状态", "登记机关", "成立日期", "营业期限",
        "企业类型", "主体类型", "经营者", "法定代表人", "经营范围",
    )
    registry_score = sum(key in compact for key in registry_fields)
    registry_context = any(key in compact for key in ("工商信息", "历史工商信息", "企查查", "天眼查", "爱企查"))
    if (registry_name and registry_score >= 2) or (registry_context and registry_score >= 2):
        return "企业信息截图"
    signatures = [
        ("申报表", ("福建省职业技能等级认定申报表", "职业技能等级认定申报表")),
        ("身份证", ("中华人民共和国居民身份证", "公民身份号码", "签发机关")),
        ("学信网学籍证明", ("教育部学籍在线验证报告", "学籍在线验证报告")),
        ("学信网学历证明", ("教育部学历证书电子注册备案表", "中国高等教育学历认证报告")),
        ("学信网学位证明", ("中国高等教育学位在线验证报告", "学位在线验证报告")),
        ("职业技能等级证书", ("职业技能等级证书", "CertificateofOccupationalSkillLevel")),
        ("工作年限承诺书", ("工作年限承诺书", "工作年限承诺函")),
        ("企业信息截图", ("统一社会信用代码", "工商注册号", "经营范围", "登记状态", "经营状态")),
        ("工作证明", ("工作证明", "兹证明", "在我单位工作")),
        ("学历证明", ("毕业证书", "毕业证明", "修业期满", "准予毕业", "毕字")),
    ]
    for label, keys in signatures:
        if any(k in compact for k in keys):
            return label
    return kind


def _docx_pages(path: Path) -> list[str]:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = []
            seen_cells = set()
            for cell in row.cells:
                # 合并单元格会由python-docx重复返回，避免重复污染列位置。
                cell_key = id(cell._tc)
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                # cell.text不包含浮动文本框；从单元格XML读取全部w:t，
                # 同时保留每个表格行和单元格的边界。
                texts = [
                    node.text for node in cell._tc.iter()
                    if node.tag.endswith("}t") and node.text
                ]
                cells.append(re.sub(r"[\r\n]+", " ", "".join(texts)).strip())
            if any(cells):
                parts.append("\t".join(cells))
    for section in doc.sections:
        parts.extend(p.text for p in section.header.paragraphs if p.text.strip())
        parts.extend(p.text for p in section.footer.paragraphs if p.text.strip())
    return ["\n".join(parts)]


def _render_pdf_page(page: fitz.Page, dpi: int) -> Image.Image:
    pix = page.get_pixmap(dpi=dpi, alpha=False)
    return Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")


def read_material(person: str, path: Path, cfg: AppConfig, ocr: LocalTesseractOCR) -> Material:
    # 文件名不参与最终材料类型判定。文档以正文/OCR内容分类，证件照以图像
    # 构图分类；文件名中即使含“身份证/学历/照片”等字样也不能覆盖内容结论。
    m = Material(person=person, path=path, document_type="其他材料")
    raster_image: Image.Image | None = None
    rendered_pdf_pages: list[tuple[int, Image.Image]] = []
    try:
        suffix = path.suffix.lower()
        if suffix == ".docx":
            m.text_pages = _docx_pages(path)
        elif suffix == ".pdf":
            doc = fitz.open(path)
            for page in doc:
                text = page.get_text("text").strip()
                if len(re.sub(r"\s", "", text)) >= 30:
                    m.text_pages.append(text)
                else:
                    image = _render_pdf_page(page, cfg.image_dpi)
                    rendered_pdf_pages.append((page.number + 1, image))
                    page_text = ocr.recognize(image)
                    m.text_pages.append(page_text)
                    if page_text and ocr.last_confidence is not None:
                        m.ocr_confidence = (
                            ocr.last_confidence if m.ocr_confidence is None
                            else min(m.ocr_confidence, ocr.last_confidence)
                        )
        else:
            raster_image = ImageOps.exif_transpose(Image.open(path)).convert("RGB")
            photo_reasons = assess_id_photo(raster_image)
            if not photo_reasons:
                m.document_type = "证件照"
            else:
                m.text_pages = [ocr.recognize(raster_image)]
                m.ocr_confidence = ocr.last_confidence

        if m.document_type != "证件照":
            m.document_type = refine_document_type("其他材料", "\n".join(m.text_pages))

        # 完成内容分类后再执行对应材料的质量门禁，避免错误文件名改变处理路径。
        if m.document_type == "证件照":
            if raster_image is not None:
                m.quality_reasons = assess_id_photo(raster_image)
            elif rendered_pdf_pages:
                for page_no, image in rendered_pdf_pages:
                    m.quality_reasons.extend(f"第{page_no}页：{reason}" for reason in assess_id_photo(image))
        elif m.document_type == "身份证":
            if raster_image is not None:
                m.quality_reasons = assess_id_image(raster_image, cfg.quality)
            else:
                for page_no, image in rendered_pdf_pages:
                    m.quality_reasons.extend(f"第{page_no}页：{reason}" for reason in assess_id_image(image, cfg.quality))
        if m.quality_reasons:
            m.quality_status = "退回"
        if m.document_type == "工作证明" and suffix in {".pdf", ".jpg", ".jpeg", ".png"}:
            stamp_found = False
            if suffix == ".pdf":
                doc = fitz.open(path)
                stamp_found = any(has_red_stamp(_render_pdf_page(page, 140)) for page in doc)
            else:
                stamp_found = has_red_stamp(raster_image or ImageOps.exif_transpose(Image.open(path)).convert("RGB"))
            if stamp_found:
                if not m.text_pages: m.text_pages = [""]
                m.text_pages[0] += "\n[检测到红色公章]"
        if (m.document_type != "证件照"
                and not any(x.strip() for x in m.text_pages)
                and not m.quality_reasons):
            m.errors.append("未提取到可用文字")
    except Exception as exc:
        m.errors.append(str(exc))
    return m
